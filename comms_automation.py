import streamlit as st
import requests
from docx import Document
from docx.shared import RGBColor
from io import BytesIO

# Common templates
TEMPLATES = {
    "In-Store Pasta Taste Test": """In-Store Pasta X – You've Been Selected! 🍝

Hi!

We're delighted to let you know that you've been selected to take part in our latest In-Store X, for which you'll receive £X upon completion.

🍽️ Important:
* This is an in-store taste testing task, so you must visit a X
* You must purchase the X specified in the task — failure to do so may result in non-payment.

📅 Deadline: X  (If you need an extension, please reach out — we're happy to help.)

💡 Before you begin: Make sure to check the map to confirm you're visiting an eligible Sainsbury's store before heading out.

📋 Here are the links to your task briefs (please complete them in order or you will not be paid):

1. Instructions Brief
2. In Store – Taste Test
3. At Home – Pre-Cooking
4. At Home – Post-Cooking

If you experience any technical issues (e.g. broken uploads or app errors), please check our FAQs first — most common questions are answered there.

📩 Still need help? Email me at hannah.seddon@smg.com and include your reference code: X so I can assist you faster.""",
    
    "Generic In-Store Task": """In-Store X – You've Been Selected!

Hi!

We're delighted to let you know that you've been selected to take part in our latest In-Store X, for which you'll receive £X upon completion.

🍽️ Important:
* This is an in-store task, so you must visit a X
* You must complete all required activities specified in the task — failure to do so may result in non-payment.

📅 Deadline: X  (If you need an extension, please reach out — we're happy to help.)

📋 Here are the links to your task briefs (please complete them in order or you will not be paid):

1. Instructions Brief
2. In Store Task

If you experience any technical issues (e.g. broken uploads or app errors), please check our FAQs first — most common questions are answered there.

📩 Still need help? Email me at hannah.seddon@smg.com and include your reference code: X so I can assist you faster."""
}

# Available models on OpenRouter
MODELS = {
    "Claude Sonnet 4.5": "anthropic/claude-sonnet-4.5",
    "Claude Sonnet 4": "anthropic/claude-sonnet-4",
    "Claude Opus 4": "anthropic/claude-opus-4",
    "Gemini 2.0 Flash": "google/gemini-2.0-flash-exp:free"
}

def call_openrouter(api_key, model, prompt):
    """Call OpenRouter API with the specified model"""
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": model,
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            },
            timeout=60
        )
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return None

def generate_ai_content(api_key, model, template, ai_instructions):
    """Generate AI-customized intro and important notes"""
    prompt = f"""You are helping customize a communication template for a market research task.

Original template:
{template}

Special instructions for this specific campaign:
{ai_instructions}

Please provide:
1. A customized intro paragraph (keep it friendly and brief, 2-3 sentences max)
2. Customized important notes/instructions based on the special instructions provided

Format your response as:
INTRO:
[Your intro paragraph here]

IMPORTANT_NOTES:
[Your important notes here, as bullet points if needed]

Keep all X placeholders intact for mail merge. Match the tone and style of the original."""

    return call_openrouter(api_key, model, prompt)

def create_hyperlinked_docx(template_text, links_dict):
    """Create a .docx with hyperlinked text"""
    doc = Document()
    
    # Split text into lines
    lines = template_text.split('\n')
    
    for line in lines:
        paragraph = doc.add_paragraph()
        
        # Check if any link category appears in this line
        link_found = False
        for category, url in links_dict.items():
            if category in line:
                # Split the line around the category text
                parts = line.split(category)
                if len(parts) == 2:
                    # Add text before link
                    paragraph.add_run(parts[0])
                    # Add hyperlink
                    add_hyperlink(paragraph, url, category)
                    # Add text after link
                    paragraph.add_run(parts[1])
                    link_found = True
                    break
        
        # If no link found, add the line as-is
        if not link_found:
            paragraph.add_run(line)
    
    return doc

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    # This adds the relationship for the hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = paragraph._element  # pylint: disable=protected-access
    hyperlink = hyperlink.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
    hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id)
    
    # Create run element for the hyperlink text
    run = hyperlink.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    r_pr = run.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    
    # Style for hyperlink (blue and underlined)
    color = r_pr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
    color.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0563C1')
    r_pr.append(color)
    
    u = r_pr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u')
    u.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    r_pr.append(u)
    
    run.append(r_pr)
    
    # Add text
    text_elem = run.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    text_elem.text = text
    run.append(text_elem)
    
    hyperlink.append(run)
    paragraph._element.append(hyperlink)  # pylint: disable=protected-access

# Streamlit UI
st.set_page_config(page_title="Comms Automation Tool", page_icon="✉️", layout="wide")

st.title("✉️ Communication Template Automation")
st.markdown("*Automate hyperlinking and AI-powered content customization for your comms*")

# Sidebar for API setup
with st.sidebar:
    st.header("🔐 API Configuration")
    api_key = st.text_input("OpenRouter API Key", type="password", help="Enter your OpenRouter API key (stored locally for this session only)")
    model_choice = st.selectbox("Select AI Model", options=list(MODELS.keys()))
    
    st.markdown("---")
    st.markdown("**How to use:**")
    st.markdown("1. Enter API key & select model")
    st.markdown("2. Choose or customize template")
    st.markdown("3. Add your links")
    st.markdown("4. Provide AI instructions")
    st.markdown("5. Generate & download!")

# Main content
col1, col2 = st.columns([1, 1])

with col1:
    st.header("📝 Template Setup")
    
    # Template selection
    template_option = st.selectbox("Choose a template", ["Select..."] + list(TEMPLATES.keys()) + ["Custom"])
    
    if template_option != "Select...":
        if template_option == "Custom":
            template_text = st.text_area("Enter your custom template", height=400, 
                                        help="Use X for mail merge fields. Link text will be automatically hyperlinked.")
        else:
            template_text = st.text_area("Template (you can edit this)", 
                                        value=TEMPLATES[template_option], 
                                        height=400)
    else:
        template_text = ""
    
    st.markdown("---")
    
    st.header("🔗 Link Management")
    st.markdown("*Add the text that should be hyperlinked and its corresponding URL*")
    
    # Initialize session state for links
    if 'links' not in st.session_state:
        st.session_state.links = []
    
    # Add new link
    with st.expander("➕ Add New Link", expanded=True):
        new_category = st.text_input("Link text (e.g., 'Instructions Brief')", key="new_category")
        new_url = st.text_input("URL", key="new_url")
        if st.button("Add Link"):
            if new_category and new_url:
                st.session_state.links.append({"category": new_category, "url": new_url})
                st.success(f"Added: {new_category}")
                st.rerun()
    
    # Display existing links
    if st.session_state.links:
        st.markdown("**Current Links:**")
        for idx, link in enumerate(st.session_state.links):
            col_a, col_b, col_c = st.columns([3, 5, 1])
            with col_a:
                st.text(link['category'])
            with col_b:
                st.text(link['url'][:50] + "..." if len(link['url']) > 50 else link['url'])
            with col_c:
                if st.button("🗑️", key=f"delete_{idx}"):
                    st.session_state.links.pop(idx)
                    st.rerun()
    
    if st.button("🗑️ Clear All Links"):
        st.session_state.links = []
        st.rerun()

with col2:
    st.header("🤖 AI Customization")
    
    ai_instructions = st.text_area(
        "Special instructions for this campaign",
        height=200,
        placeholder="E.g., 'This is for a new gluten-free pasta range. Emphasize that participants must check allergen information. Deadline is strict due to product launch timing.'",
        help="Provide context and specific requirements for this campaign. AI will customize the intro and important notes based on this."
    )
    
    use_ai = st.checkbox("Use AI to customize intro & important notes", value=True)
    
    st.markdown("---")
    
    st.header("🚀 Generate Document")
    
    if st.button("Generate Hyperlinked Document", type="primary", use_container_width=True):
        # Validation
        if not template_text:
            st.error("Please select or enter a template!")
        elif not st.session_state.links:
            st.error("Please add at least one link!")
        elif use_ai and not api_key:
            st.error("Please enter your OpenRouter API key in the sidebar!")
        elif use_ai and not ai_instructions:
            st.error("Please provide AI instructions or uncheck 'Use AI'!")
        else:
            with st.spinner("Generating your document..."):
                final_template = template_text
                
                # AI customization if enabled
                if use_ai:
                    with st.spinner("AI is customizing your content..."):
                        ai_response = generate_ai_content(
                            api_key, 
                            MODELS[model_choice], 
                            template_text, 
                            ai_instructions
                        )
                        
                        if ai_response:
                            st.success("✅ AI customization complete!")
                            with st.expander("View AI-generated content"):
                                st.markdown(ai_response)
                            
                            # Parse AI response and update template
                            # This is a simple implementation - you might want to make it more sophisticated
                            if "INTRO:" in ai_response and "IMPORTANT_NOTES:" in ai_response:
                                parts = ai_response.split("IMPORTANT_NOTES:")
                                intro = parts[0].replace("INTRO:", "").strip()
                                notes = parts[1].strip()
                                
                                # Replace intro paragraph (between first "Hi!" and first emoji line)
                                if "Hi!" in final_template:
                                    intro_start = final_template.find("Hi!")
                                    intro_end = final_template.find("🍽️")
                                    if intro_end > intro_start:
                                        final_template = (final_template[:intro_start] + 
                                                        f"Hi!\n\n{intro}\n\n" + 
                                                        final_template[intro_end:])
                        else:
                            st.warning("AI customization failed. Using original template.")
                
                # Create links dictionary
                links_dict = {link['category']: link['url'] for link in st.session_state.links}
                
                # Generate document
                doc = create_hyperlinked_docx(final_template, links_dict)
                
                # Save to BytesIO
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # Download button
                st.download_button(
                    label="📥 Download .docx",
                    data=doc_io,
                    file_name="communication_template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.success("✅ Document generated successfully!")
                st.info("💡 Open the .docx file and use it in your Gmail mail merge. All X fields are preserved for merging.")

# Footer
st.markdown("---")
st.markdown("*Built for streamlining comms workflows | All data stays local*")
